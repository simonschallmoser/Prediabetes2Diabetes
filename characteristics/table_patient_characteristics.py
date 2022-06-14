import numpy as np
import pandas as pd
import pickle
from docx import Document
from docx.shared import Cm
from scipy.stats import ttest_ind

from patient_characteristics import demo_characteristics, lab_characteristics, icd_characteristics, med_characteristics
from categories_dicts import med_to_class, drug_class_to_name, rename_tests, ordered_tests
import data_reader
import preprocess
import helper
import icd9

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                                    
def table_patient_characteristics(df_lab, 
                                  df_icd, 
                                  df_demo, 
                                  df_med, 
                                  antidiabetes_medication, 
                                  long,
                                  style, 
                                  col_width, 
                                  table_name='patient_characteristics', 
                                  both=False, 
                                  p_values=False,
                                  differentiate_diabetes_patients=True):
    
    """
    Generates table (Table S1) with patient characteristics including all predictors if long is True, 
    otherwise it generates short table (Table 1 in main text).
    """

    # Initizalize table
    
    doc = Document()
    table = doc.add_table(0, 0)
    table.style = style
    
    if both:
        for i in range(12):
            table.add_column(Cm(col_width))
        if p_values:
            table.add_column(Cm(col_width))
    else:
        if differentiate_diabetes_patients:
            for i in range(10):
                table.add_column(Cm(col_width))
            if p_values:
                table.add_column(Cm(col_width))
        else:
            for i in range(8):
                table.add_column(Cm(col_width))
            if p_values:
                table.add_column(Cm(col_width))
    
    ids_diabetes_icd_med, ids_diabetes_lab = helper.compute_ids_diabetes_onset(df_lab,
                                                                               df_icd,
                                                                               df_med,
                                                                               antidiabetes_medication,
                                                                               return_diagnosis_criteria_indices=True)
    ids_prediabetes_onset, ids_diabetes_onset = helper.ids(df_lab, df_icd, df_med, df_demo, antidiabetes_medication)
    data_prep = preprocess.preprocess_data(df_lab, df_demo, ids_prediabetes_onset, ids_diabetes_onset, df_icd, df_med,
                                           ffill_lab=False, ffill_demo=False)
    ids_diabetes_icd_med = [i for i in ids_diabetes_icd_med if i in ids_diabetes_onset.keys()]
    ids_diabetes_lab = [i for i in ids_diabetes_lab if i in ids_diabetes_onset.keys()]
    # Define index for overall diabetes column
    
    if p_values:
        idx_diab_all = -2
    else:
        idx_diab_all = -1

    tests = [i.split('=')[1] for i in data_prep.columns if 'test' in i]
    icds = [i.split('=')[1].split('_')[0] for i in data_prep.columns if 'icd' in i]
    med_cols = [i for i in data_prep.columns if 'med' in i]
    try:
        med_cols.remove('(med_class = 20)')
    except:
        pass
    try:
        med_cols.remove('(med_class = 21)')
    except:
        pass
    try:
        med_cols.remove('(med_class = 22)')
    except:
        pass        
        
    demo_dicts = []
    lab_dicts = []
    icd_dicts = []
    med_dicts = []
    ids_pre_list = []
    ids_diab_list = []
    for pred_year in range(1, 6):
        print('Forecast horizon:', pred_year)
        data = preprocess.final_preprocessing_step(data_prep, ids_diabetes_onset, pred_year)
        data_ids = list(data.index.get_level_values('id').unique())
        ids_pre_list.append(list(data[data['y'] == 0].index.get_level_values('id').unique()))
        ids_diab_list.append(list(data[data['y'] == 1].index.get_level_values('id').unique()))
        
        demo_dicts.append(demo_characteristics(data_prep, ids_prediabetes_onset, ids_diabetes_onset,
                                               ids_diabetes_icd_med, ids_diabetes_lab, pred_year))
        lab_dicts.append(lab_characteristics(data_prep, ids_prediabetes_onset, ids_diabetes_onset, 
                                             ids_diabetes_icd_med, ids_diabetes_lab, pred_year, long))
        icd_dicts.append(icd_characteristics(data_prep, ids_prediabetes_onset, ids_diabetes_onset, 
                                             ids_diabetes_icd_med, ids_diabetes_lab, pred_year))
        med_dicts.append(med_characteristics(data_prep, ids_prediabetes_onset, ids_diabetes_onset, 
                                             ids_diabetes_icd_med, ids_diabetes_lab, pred_year))
    
    table.add_row()
    row = table.rows[0]
    if both:
        for i in range(5):
            row.cells[2*i+1].text = 'Prediabetes'
            row.cells[2*i+2].text = 'Diabetes'
            row.cells[idx_diab_all].text = 'Characteristics at diabetes diagnosis'
            if p_values:
                row.cells[-1].text = 'p-value'
    else:
        row.cells[4].text = 'Diabetes'
        row.cells[1].text = 'Prediabetes'
        row.cells[idx_diab_all].text = 'Characteristics at diabetes diagnosis'
        if p_values:
            row.cells[-1].text = 'p-value'
    table.add_row()
    row = table.rows[len(table.rows)-1]
    row.cells[-3].text = 'Diagnosis criteria: ICD-9 codes and antidiabetes medication'
    row.cells[-2].text = 'Diagnosis criterium: HbA1c'
    row.cells[-1].text = 'All'
    table.add_row()
    row = table.rows[len(table.rows)-1]
    row.cells[0].text = 'Forecast horizon [years]'
    for i in range(1, 6):
        if both:
            row.cells[2*i].text = str(i)
        else:
            row.cells[i+1].text = str(i)
    table.add_row()
    row = table.rows[len(table.rows)-1]
    row.cells[0].text = 'Number of samples'
    if both:
        for i in range(5):
            row.cells[2*i+1].text = str(len(ids_pre_list[i]))
            row.cells[2*i+2].text = str(len(ids_diab_list[i]))
        row.cells[idx_diab_all].text = str(len(ids_diab_list[-1]))
    else:
        row.cells[1].text = str(len(ids_pre_list[0])+len(ids_diab_list[0]))
        for i in range(5):
            row.cells[i+2].text = str(len(ids_diab_list[i]))
        row.cells[idx_diab_all].text = str(len(ids_diab_list[-1]))
        if differentiate_diabetes_patients:
            row.cells[-3].text = str(len(ids_diabetes_icd_med))
            row.cells[-2].text = str(len(ids_diabetes_lab))
    table.add_row()
    row = table.rows[len(table.rows)-1]
    row.cells[0].text = 'Incidence [%]'
    for i in range(5):
        if both:
            row.cells[2*i+1].text = str(np.round(len(ids_diab_list[i])/len(ids_prediabetes_onset)*100, 1))
        else:
            row.cells[i+2].text = str(np.round(len(ids_diab_list[i])/len(ids_prediabetes_onset)*100, 1))
    row.cells[idx_diab_all].text = str(np.round(len(ids_diab_list[-1])/len(ids_prediabetes_onset)*100, 1))
    if differentiate_diabetes_patients:
        row.cells[-3].text = str(np.round(len(ids_diabetes_icd_med)/len(ids_prediabetes_onset)*100, 1))
        row.cells[-2].text = str(np.round(len(ids_diabetes_lab)/len(ids_prediabetes_onset)*100, 1))
    row.cells[1].text = '-'
    table.add_row()
    row = table.rows[len(table.rows)-1]
    row.cells[0].text = 'Demographic data'
    make_rows_bold(table.rows[0])
    make_rows_bold(table.rows[1])
    make_rows_bold(table.rows[5])
    table.add_row()
    row_sex = table.rows[len(table.rows)-1]
    row_sex.cells[0].text = 'Sex'
    table.add_row()
    row_male = table.rows[len(table.rows)-1]
    row_male.cells[0].text = '    Male'
    table.add_row()
    row_female = table.rows[len(table.rows)-1]
    row_female.cells[0].text = '    Female'
    row_length = len(table.rows)
    table.add_row()
    row_age_pred = table.rows[len(table.rows)-1]
    row_age_pred.cells[0].text = 'Age [years]'
    table.add_row()
    row_bmi = table.rows[len(table.rows)-1]
    row_bmi.cells[0].text = 'BMI [kg/m^2]'
    if long:
        table.add_row()
        row_sys = table.rows[len(table.rows)-1]
        row_sys.cells[0].text = 'Systolic blood pressure [mmHg]'
        table.add_row()
        row_dia = table.rows[len(table.rows)-1]
        row_dia.cells[0].text = 'Diastolic blood pressure [mmHg]'
    table.add_row()
    row_bio = table.rows[len(table.rows)-1]
    row_bio.cells[0].text = 'Biomarkers'
    make_rows_bold(table.rows[len(table.rows)-1])

    for idx, demo_dict in enumerate(demo_dicts):
        
        if both:
            cell_diab_idx = 2*idx+2
            cell_pre_idx = 2*idx+1
        else:
            cell_diab_idx = idx+2
            cell_pre_idx = 1
        row_male.cells[cell_diab_idx].text = str(demo_dict['males_diab']) + ' (' \
        + str(demo_dict['males_diab_pct']) + '%)'
        row_female.cells[cell_diab_idx].text = str(demo_dict['females_diab']) + ' (' \
        + str(demo_dict['females_diab_pct']) + '%)'
        row_age_pred.cells[cell_diab_idx].text = str(demo_dict['age_diab_mean']) + ' (' \
        + str(demo_dict['age_diab_std']) + ')'
        row_bmi.cells[cell_diab_idx].text = str(demo_dict['bmi_diab_mean']) + ' (' \
        + str(demo_dict['bmi_diab_std']) + ')'
        if long:
            row_sys.cells[cell_diab_idx].text = str(demo_dict['systolic_bp_diab_mean']) + ' (' \
            + str(demo_dict['systolic_bp_diab_std']) + ')'
            row_dia.cells[cell_diab_idx].text = str(demo_dict['diastolic_bp_diab_mean']) + ' (' \
            + str(demo_dict['diastolic_bp_diab_std']) + ')'
        
        row_male.cells[cell_pre_idx].text = str(demo_dict['males_pre']) + ' (' \
        + str(demo_dict['males_pre_pct']) + '%)'
        row_female.cells[cell_pre_idx].text = str(demo_dict['females_pre']) + ' (' \
        + str(demo_dict['females_pre_pct']) + '%)'
        row_age_pred.cells[cell_pre_idx].text = str(demo_dict['age_pre_mean']) + ' (' \
        + str(demo_dict['age_pre_std']) + ')'
        row_bmi.cells[cell_pre_idx].text = str(demo_dict['bmi_pre_mean']) + ' (' \
        + str(demo_dict['bmi_pre_std']) + ')'
        if long:
            row_sys.cells[cell_pre_idx].text = str(demo_dict['systolic_bp_pre_mean']) + ' (' \
            + str(demo_dict['systolic_bp_pre_std']) + ')'
            row_dia.cells[cell_pre_idx].text = str(demo_dict['diastolic_bp_pre_mean']) + ' (' \
            + str(demo_dict['diastolic_bp_pre_std']) + ')'
    
    if differentiate_diabetes_patients:
        idx_list = [-3, -2, -1]
        data_type_list = ['diab_all_icd_med', 'diab_all_lab', 'diab_all']
    else:
        idx_list = [-1]
        data_type_list = ['diab_all']

    
    for idx, data_type in zip(idx_list, data_type_list):
    
        row_male.cells[idx].text = str(demo_dict[f'males_{data_type}']) + ' (' \
        + str(demo_dict[f'males_{data_type}_pct']) + '%)'
        row_female.cells[idx].text = str(demo_dict[f'females_{data_type}']) + ' (' \
        + str(demo_dict[f'females_{data_type}_pct']) + '%)'

        row_age_pred.cells[idx].text = str(demo_dict[f'age_{data_type}_mean']) + ' (' \
        + str(demo_dict[f'age_{data_type}_std']) + ')'
        row_bmi.cells[idx].text = str(demo_dict[f'bmi_{data_type}_mean']) + ' (' \
        + str(demo_dict[f'bmi_{data_type}_std']) + ')'
        if long:
            row_sys.cells[idx].text = str(demo_dict[f'systolic_bp_{data_type}_mean']) + ' (' \
            + str(demo_dict[f'systolic_bp_{data_type}_std']) + ')'
            row_dia.cells[idx].text = str(demo_dict[f'diastolic_bp_{data_type}_mean']) + ' (' \
            + str(demo_dict[f'diastolic_bp_{data_type}_std']) + ')'
    
    if p_values:
        p_values = []
        p_values.append(ttest_ind(demo_dict['age_pre_dist'], 
                                  demo_dict['age_diab_all_dist'], nan_policy='omit').pvalue)
        p_values.append(ttest_ind(demo_dict['bmi_pre_dist'], 
                                  demo_dict['bmi_diab_all_dist'], nan_policy='omit').pvalue)
        p_values.append(ttest_ind(demo_dict['systolic_bp_pre_dist'],
                                  demo_dict['systolic_bp_diab_all_dist'], nan_policy='omit').pvalue)
        p_values.append(ttest_ind(demo_dict['diastolic_bp_pre_dist'],
                                  demo_dict['diastolic_bp_diab_all_dist'], nan_policy='omit').pvalue)

        for row_index, p_value in enumerate(p_values):
            row = table.rows[row_length+row_index]
            if p_value < 0.001:
                row.cells[-1].text = '<0.001'
            elif p_value < 0.01:
                row.cells[-1].text = '<0.01'
            elif p_value < 0.05:
                row.cells[-1].text = '<0.05'
            else:
                row.cells[-1].text = str(np.round(p_value, 2))
    
    if not long:
        ordered_tests1 = ['test=glucose', 'test=hba1c(%)', 'test=hba1c(mmol/mol)_1']
    else:
        ordered_tests1 = ordered_tests
    row_length = len(table.rows)
    for row_index, test in enumerate(ordered_tests1):
        test_name = test.split('=')[1]
        table.add_row()
        row = table.rows[row_length+row_index]
        row.cells[0].text = rename_tests[test]
        try:
            for idx, lab_dict in enumerate(lab_dicts):
                if both:
                    cell_diab_idx = 2*idx+2
                    cell_pre_idx = 2*idx+1
                else:
                    cell_diab_idx = idx+2
                    cell_pre_idx = 1
                row.cells[cell_diab_idx].text = str(lab_dict[test_name+'_diab_mean']) \
                + ' (' + str(lab_dict[test_name+'_diab_std']) + ')'
                row.cells[cell_pre_idx].text = str(lab_dict[test_name+'_pre_mean']) \
                + ' (' + str(lab_dict[test_name+'_pre_std']) + ')'
            
            if differentiate_diabetes_patients:
                idx_list = [-3, -2, -1]
                data_type_list = ['diab_all_icd_med', 'diab_all_lab', 'diab_all']
            else:
                idx_list = [-1]
                data_type_list = ['diab_all']


            for idx, data_type in zip(idx_list, data_type_list):
                row.cells[idx].text = str(lab_dict[test_name+'_'+data_type+'_mean']) \
                + ' (' + str(lab_dict[test_name+'_'+data_type+'_std']) + ')'
                
            if p_values:
                p_value = ttest_ind(lab_dict[test_name+'_pre_dist'],
                                    lab_dict[test_name+'_diab_all_dist'], nan_policy='omit').pvalue
                if p_value < 0.001:
                    row.cells[-1].text = '<0.001'
                elif p_value < 0.01:
                    row.cells[-1].text = '<0.01'
                elif p_value < 0.05:
                    row.cells[-1].text = '<0.05'
                else:
                    row.cells[-1].text = str(np.round(p_value, 2))
        except:
            continue

    table.add_row()
    row_length = len(table.rows)
    row = table.rows[row_length-1]
    row.cells[0].text = 'Comorbidities'
    if not long:
        icds = ['272', '401']
    make_rows_bold(table.rows[row_length-1])
    for row_index, icd_code in enumerate(icds):
        table.add_row()
        row = table.rows[row_length+row_index]
        if icd_code == '272':
            row.cells[0].text = 'Dyslipidemia'
        elif icd_code == '401':
            row.cells[0].text = 'Hypertension'
        else:
            row.cells[0].text = icd9.icd9_to_description[int(icd_code)]
        for idx, icd_dict in enumerate(icd_dicts):
            if both:
                cell_diab_idx = 2*idx+2
                cell_pre_idx = 2*idx+1
            else:
                cell_diab_idx = idx+2
                cell_pre_idx = 1
            row.cells[cell_diab_idx].text = str(int(icd_dict[str(icd_code)+'_diab'])) \
            + ' (' + str(icd_dict[str(icd_code)+'_diab_pct']) + '%)'
            row.cells[cell_pre_idx].text = str(int(icd_dict[str(icd_code)+'_pre'])) \
            + ' (' + str(icd_dict[str(icd_code)+'_pre_pct']) + '%)'
            
        if differentiate_diabetes_patients:
            idx_list = [-3, -2, -1]
            data_type_list = ['diab_all_icd_med', 'diab_all_lab', 'diab_all']
        else:
            idx_list = [-1]
            data_type_list = ['diab_all']


        for idx, data_type in zip(idx_list, data_type_list):
            row.cells[idx].text = str(int(icd_dict[str(icd_code)+'_'+data_type])) \
            + ' (' + str(icd_dict[str(icd_code)+'_'+data_type+'_pct']) + '%)'
            
        if p_values:
            row.cells[-1].text = '-'
    
    if long:
        table.add_row()
        row_length = len(table.rows)
        row = table.rows[row_length-1]
        row.cells[0].text = 'Medications'
        make_rows_bold(table.rows[row_length-1])
        for row_index, med_col in enumerate(med_cols):
            med_class = int(med_col.split('= ')[1].split(')')[0])
            table.add_row()
            row = table.rows[row_length+row_index]
            row.cells[0].text = drug_class_to_name[med_class]
            for idx, med_dict in enumerate(med_dicts):
                if both:
                    cell_diab_idx = 2*idx+2
                    cell_pre_idx = 2*idx+1
                else:
                    cell_diab_idx = idx+2
                    cell_pre_idx = 1
                row.cells[cell_diab_idx].text = str(int(med_dict[str(med_class)+'_diab'])) \
                + ' (' +str(med_dict[str(med_class)+'_diab_pct']) + '%)'
                row.cells[cell_pre_idx].text = str(int(med_dict[str(med_class)+'_pre'])) \
                + ' (' + str(med_dict[str(med_class)+'_pre_pct']) + '%)'

            if differentiate_diabetes_patients:
                idx_list = [-3, -2, -1]
                data_type_list = ['diab_all_icd_med', 'diab_all_lab', 'diab_all']
            else:
                idx_list = [-1]
                data_type_list = ['diab_all']


            for idx, data_type in zip(idx_list, data_type_list):
                row.cells[idx].text = str(int(med_dict[str(med_class)+'_'+data_type])) \
                + ' (' +str(med_dict[str(med_class)+'_'+data_type+'_pct']) + '%)'

            if p_values:
                row.cells[-1].text = '-'
    
    if long:
        add = 'long'
    else:
        add = 'short'
    doc.save(f'{table_name}_{add}.docx')